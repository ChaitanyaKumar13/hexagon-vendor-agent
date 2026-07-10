"""
Hexagon Vendor Registration Agent
=================================
Enter-once vendor onboarding: one guided form feeds every registration
document, with deterministic Indian statutory validation (PAN / GSTIN /
IFSC / TAN / Udyam, including the GSTIN-embeds-PAN cross-check), AI
certificate verification (Claude vision), generated ready-to-sign
documents (Annex 4 + 206AB declaration), and a password-protected admin
review dashboard.

Storage:  Google Sheets (same service account as the service centre app;
          a separate spreadsheet).
AI layer: Claude - certificate reading & cross-check only. All acceptance
          decisions stay with humans. Exhibit 7 (anti-corruption
          questionnaire) is deliberately NOT auto-filled.
"""

from __future__ import annotations

import base64
import io
import json
import re
import time
import uuid
from datetime import date, datetime
from pathlib import Path

import gspread
import pandas as pd
import streamlit as st
from google.oauth2.service_account import Credentials

# ---------------------------------------------------------------------------
# Branding
# ---------------------------------------------------------------------------

BRANDING = {
    "navy_bg": "#0C2C40",
    "navy_panel": "#123B54",
    "lime": "#C9DD28",
    "cyan": "#6FD6FF",
    "accent": "#0096D6",
}

st.set_page_config(page_title="Hexagon Vendor Registration",
                   page_icon="⬡", layout="wide")

BG, PANEL = BRANDING["navy_bg"], BRANDING["navy_panel"]
LIME, CYAN = BRANDING["lime"], BRANDING["cyan"]
INPUT_BG, TEXT = "#1D465E", "#D7E3EC"

st.markdown(
    f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Hanken+Grotesk:wght@300;400;500;600;700;800&display=swap');
    html, body, [class*="css"], .stApp, p, li, label, input, textarea, button, select {{
        font-family: 'Hanken Grotesk', sans-serif !important;
    }}
    [data-testid="stIconMaterial"], .material-symbols-rounded, .material-symbols-outlined {{
        font-family: 'Material Symbols Rounded', 'Material Symbols Outlined' !important;
    }}
    .stApp {{
        background:
            radial-gradient(1100px 500px at 80% -10%, rgba(0,150,214,0.25), transparent 60%),
            radial-gradient(900px 500px at -10% 110%, rgba(201,221,40,0.10), transparent 55%),
            {BG};
        color: {TEXT};
    }}
    .stApp::before {{
        content: ""; position: fixed; top: 0; left: 0; right: 0; height: 4px; z-index: 1000;
        background: linear-gradient(90deg, {LIME}, {CYAN}, {LIME});
        background-size: 200% 100%; animation: hexbar 6s linear infinite;
    }}
    @keyframes hexbar {{ 0% {{background-position: 0% 0;}} 100% {{background-position: 200% 0;}} }}
    h1, h2, h3, h4 {{ color: #FFFFFF !important; }}
    .hex-title {{
        font-size: 2rem; font-weight: 800;
        background: linear-gradient(90deg, #FFFFFF 30%, {CYAN});
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    }}
    .hex-sub {{ color: {TEXT}; opacity: 0.85; margin-top: 0.2rem; }}
    .stTextInput input, .stTextArea textarea, [data-baseweb="input"] input,
    [data-baseweb="base-input"] input {{
        color: #FFFFFF !important; -webkit-text-fill-color: #FFFFFF !important;
        background: {INPUT_BG} !important; background-color: {INPUT_BG} !important;
        caret-color: {LIME}; min-height: 2.6rem;
    }}
    [data-baseweb="input"], [data-baseweb="input"] > div {{ background: {INPUT_BG} !important; }}
    .stTextInput input::placeholder, .stTextArea textarea::placeholder {{
        color: rgba(215,227,236,0.55) !important;
        -webkit-text-fill-color: rgba(215,227,236,0.55) !important;
    }}
    [data-baseweb="select"] * {{ color: #FFFFFF !important; -webkit-text-fill-color: #FFFFFF !important; }}
    [data-baseweb="select"] > div {{ background: {INPUT_BG} !important; }}
    [data-baseweb="popover"] [role="listbox"], [data-baseweb="menu"] {{ background: {PANEL} !important; }}
    [data-baseweb="popover"] [role="option"] {{ color: #FFFFFF !important; }}
    [data-baseweb="popover"] [role="option"]:hover,
    [data-baseweb="popover"] [role="option"][aria-selected="true"] {{
        background: rgba(201,221,40,0.16) !important;
    }}
    .stButton > button {{
        background: {LIME}; color: {BG} !important; font-weight: 700;
        border: none; border-radius: 8px;
    }}
    .stButton > button p {{ color: {BG} !important; }}
    .stButton > button:hover {{ background: {CYAN}; }}
    #MainMenu, footer, [data-testid="stToolbar"] {{ visibility: hidden; }}
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

FIRM_TYPES = ["Individual", "Proprietor", "Partnership Firm", "LLP",
              "Private Limited / Corporate", "Other"]

SERVICE_CATEGORIES = [
    "Supply of Goods", "Professional / Consultancy Services",
    "Works Contract / Installation", "Transport & Logistics",
    "IT / Software Services", "Facility / Manpower Services",
    "Calibration / Repair Services", "Other",
]

DOC_TYPES = {
    "pan_card": "PAN Card copy",
    "gst_cert": "GST Registration Certificate",
    "cheque": "Cancelled Cheque",
    "msme_cert": "MSME / Udyam Certificate (if applicable)",
    "incorporation": "Incorporation Certificate / Partnership Deed / LLP Agreement",
}

CLAUDE_MODEL = "claude-sonnet-5"

SHEET_COLUMNS = [
    "vendor_id", "status", "created_at", "updated_at",
    "legal_name", "trade_name", "registered_address", "firm_type",
    "firm_reg_no", "linked_to_government", "country_of_origin",
    "contact_person", "contact_mobile", "contact_email",
    "nature_of_service", "service_category",
    "pan", "tan", "gstin", "msme_registered", "msme_number", "esic_number",
    "itr_fy_2023_24", "itr_fy_2024_25", "itr_fy_2025_26_will_file",
    "bank_account_name", "bank_account_number", "ifsc", "swift",
    "bank_name", "bank_branch_address",
    "docs_provided", "ai_verification", "admin_notes",
]

STATUSES = ["Submitted", "Under Review", "Approved", "Rejected"]

# ---------------------------------------------------------------------------
# Validation - Indian statutory identifiers
# ---------------------------------------------------------------------------

PAN_RE = re.compile(r"^[A-Z]{5}[0-9]{4}[A-Z]$")
GSTIN_RE = re.compile(r"^[0-9]{2}[A-Z]{5}[0-9]{4}[A-Z][1-9A-Z]Z[0-9A-Z]$")
IFSC_RE = re.compile(r"^[A-Z]{4}0[A-Z0-9]{6}$")
TAN_RE = re.compile(r"^[A-Z]{4}[0-9]{5}[A-Z]$")
UDYAM_RE = re.compile(r"^UDYAM-[A-Z]{2}-\d{2}-\d{7}$")
EMAIL_RE = re.compile(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$")


def clean_upper(s: str) -> str:
    return re.sub(r"\s+", "", str(s or "")).upper()


def valid_mobile(m: str) -> bool:
    m = re.sub(r"\D", "", m)
    return len(m) == 10 and m[0] in "6789"


def validate_vendor(v: dict) -> list[str]:
    """Deterministic validation. Returns a list of problems (empty = OK)."""
    p: list[str] = []
    if not v["legal_name"].strip():
        p.append("Legal name is required.")
    if not v["registered_address"].strip():
        p.append("Registered address is required.")
    if not v["contact_person"].strip():
        p.append("Contact person name is required.")
    if not valid_mobile(v["contact_mobile"]):
        p.append("A valid 10-digit Indian mobile number is required.")
    if not EMAIL_RE.match(v["contact_email"].strip()):
        p.append("A valid contact email is required.")
    if not v["nature_of_service"].strip():
        p.append("Nature of service is required (used for TDS classification).")

    pan = clean_upper(v["pan"])
    if not PAN_RE.match(pan):
        p.append("PAN format is invalid - expected AAAAA9999A (5 letters, 4 digits, 1 letter).")

    gstin = clean_upper(v["gstin"])
    if gstin:
        if not GSTIN_RE.match(gstin):
            p.append("GSTIN format is invalid - expected 15 characters like 06AAAAA9999A1Z5.")
        elif PAN_RE.match(pan) and gstin[2:12] != pan:
            p.append(f"GSTIN and PAN don't match - the GSTIN contains PAN '{gstin[2:12]}' "
                     f"but you entered PAN '{pan}'. One of them is wrong.")

    tan = clean_upper(v["tan"])
    if tan and not TAN_RE.match(tan):
        p.append("TAN format is invalid - expected AAAA99999A.")

    if v["msme_registered"] == "Yes":
        ud = clean_upper(v["msme_number"]).replace("UDYAM", "UDYAM")  # normalise spacing already done
        if ud and not UDYAM_RE.match(v["msme_number"].strip().upper()):
            p.append("Udyam number format looks wrong - expected UDYAM-XX-00-0000000.")

    acct = re.sub(r"\D", "", v["bank_account_number"])
    if not (9 <= len(acct) <= 18):
        p.append("Bank account number should be 9-18 digits.")
    if not v["bank_account_name"].strip():
        p.append("Account holder name (as per bank records) is required.")
    if not IFSC_RE.match(clean_upper(v["ifsc"])):
        p.append("IFSC format is invalid - expected AAAA0999999 (5th character is zero).")
    if not v["bank_name"].strip():
        p.append("Bank name is required.")
    return p


# ---------------------------------------------------------------------------
# Google Sheets storage
# ---------------------------------------------------------------------------

SCOPES = ["https://www.googleapis.com/auth/spreadsheets",
          "https://www.googleapis.com/auth/drive"]


@st.cache_resource(show_spinner=False)
def get_worksheet():
    creds = Credentials.from_service_account_info(
        dict(st.secrets["gcp_service_account"]), scopes=SCOPES)
    gc = gspread.authorize(creds)
    sh = gc.open(st.secrets.get("VENDOR_SHEET_NAME", "Hexagon Vendor Registrations"))
    ws = sh.sheet1
    first = ws.row_values(1)
    if first != SHEET_COLUMNS:
        if not any(first):
            ws.update("A1", [SHEET_COLUMNS])
        else:
            raise RuntimeError("Sheet header mismatch - clear the sheet or fix the header.")
    return ws


@st.cache_data(ttl=20, show_spinner=False)
def load_vendors() -> pd.DataFrame:
    ws = get_worksheet()
    records = ws.get_all_records()
    if not records:
        return pd.DataFrame(columns=SHEET_COLUMNS)
    df = pd.DataFrame(records)
    for c in SHEET_COLUMNS:
        if c not in df.columns:
            df[c] = ""
    return df[SHEET_COLUMNS].astype(str)


def refresh():
    load_vendors.clear()


def append_vendor(row: dict):
    ws = get_worksheet()
    ws.append_row([row.get(c, "") for c in SHEET_COLUMNS],
                  value_input_option="USER_ENTERED")
    refresh()


def update_vendor(vendor_id: str, changes: dict) -> bool:
    ws = get_worksheet()
    for i, rec in enumerate(ws.get_all_records()):
        if str(rec.get("vendor_id", "")).strip().upper() == vendor_id.strip().upper():
            merged = {**rec, **changes,
                      "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M")}
            end_col = chr(ord("A") + len(SHEET_COLUMNS) - 1) if len(SHEET_COLUMNS) <= 26 \
                else "A" + chr(ord("A") + len(SHEET_COLUMNS) - 27)
            ws.update(f"A{i + 2}:{end_col}{i + 2}",
                      [[merged.get(c, "") for c in SHEET_COLUMNS]],
                      value_input_option="USER_ENTERED")
            refresh()
            return True
    return False


def gen_vendor_id(df: pd.DataFrame) -> str:
    existing = set(df["vendor_id"].str.upper()) if not df.empty else set()
    while True:
        cand = "VND-" + uuid.uuid4().hex[:6].upper()
        if cand not in existing:
            return cand


# ---------------------------------------------------------------------------
# AI certificate verification (Claude vision)
# ---------------------------------------------------------------------------

def get_claude():
    try:
        import anthropic
        return anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])
    except Exception:
        return None


VERIFY_PROMPTS = {
    "pan_card": ("a PAN card. Extract: the PAN number and the name printed on the card. "
                 "Compare with declared PAN '{pan}' and legal/contact name '{legal_name}'."),
    "gst_cert": ("a GST Registration Certificate. Extract: GSTIN, legal name, and trade name. "
                 "Compare with declared GSTIN '{gstin}' and legal name '{legal_name}'."),
    "cheque": ("a cancelled cheque. Extract: account number, IFSC code, and account holder "
               "name. Compare with declared account '{bank_account_number}', IFSC '{ifsc}', "
               "and account name '{bank_account_name}'."),
    "msme_cert": ("an MSME / Udyam certificate. Extract: Udyam registration number and "
                  "enterprise name. Compare with declared Udyam number '{msme_number}' and "
                  "legal name '{legal_name}'."),
    "incorporation": ("an incorporation certificate, partnership deed, or LLP agreement. "
                      "Extract: entity name and registration/CIN number. Compare with "
                      "declared legal name '{legal_name}' and registration no. '{firm_reg_no}'."),
}


def verify_document(client, doc_key: str, file_bytes: bytes, mime: str,
                    vendor: dict) -> dict:
    """Ask Claude to read a certificate and cross-check it against typed data.
    Read-only: the verdict is advisory, humans decide."""
    if len(file_bytes) > 4_500_000:
        return {"error": "File too large for AI verification (keep under ~4.5 MB)."}
    b64 = base64.b64encode(file_bytes).decode()
    if mime == "application/pdf":
        media_block = {"type": "document",
                       "source": {"type": "base64", "media_type": mime, "data": b64}}
    else:
        media_block = {"type": "image",
                       "source": {"type": "base64", "media_type": mime, "data": b64}}
    task = VERIFY_PROMPTS[doc_key].format(**{k: vendor.get(k, "") for k in
                                             ("pan", "gstin", "legal_name",
                                              "bank_account_number", "ifsc",
                                              "bank_account_name", "msme_number",
                                              "firm_reg_no")})
    prompt = (
        f"This document should be {task}\n\n"
        "Rules: read ONLY what is visible - never guess or fill gaps. If a value "
        "is unreadable, use null and add an issue. Name matching should ignore "
        "case, punctuation and honorifics but flag genuinely different names.\n\n"
        "Respond with ONLY a JSON object, no markdown fences, in this shape:\n"
        '{"document_looks_like": "...", "extracted": {...}, '
        '"matches": {"<field>": true/false}, "issues": ["..."], '
        '"confidence": "high|medium|low"}'
    )
    for attempt in range(2):
        try:
            resp = get_claude().messages.create(
                model=CLAUDE_MODEL, max_tokens=1500,
                messages=[{"role": "user",
                           "content": [media_block, {"type": "text", "text": prompt}]}])
            text = "".join(b.text for b in resp.content
                           if getattr(b, "type", "") == "text").strip()
            if not text:
                time.sleep(1)
                continue
            text = re.sub(r"^```(json)?|```$", "", text.strip(), flags=re.M).strip()
            return json.loads(text)
        except json.JSONDecodeError:
            return {"error": "AI returned an unreadable verdict - try again."}
        except Exception as e:  # noqa: BLE001
            return {"error": f"Verification failed ({type(e).__name__})."}
    return {"error": "AI verification produced no output - try again."}


def render_verdict(label: str, verdict: dict):
    if "error" in verdict:
        st.warning(f"**{label}:** {verdict['error']}")
        return
    matches = verdict.get("matches", {})
    issues = verdict.get("issues", [])
    all_ok = matches and all(matches.values()) and not issues
    icon = "✅" if all_ok else "⚠️"
    lines = [f"{icon} **{label}** - confidence: {verdict.get('confidence', '?')}"]
    for field, ok in matches.items():
        lines.append(f"- {field}: {'match ✓' if ok else '**MISMATCH ✗**'}")
    for iss in issues:
        lines.append(f"- ⚠️ {iss}")
    st.markdown("\n".join(lines))


# ---------------------------------------------------------------------------
# Generated documents (python-docx) - ready-to-sign
# ---------------------------------------------------------------------------

def _docx_form(title: str, rows: list[tuple[str, str]], footer_lines: list[str]) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    h = doc.add_heading(title, level=1)
    for label, value in rows:
        para = doc.add_paragraph()
        run = para.add_run(f"{label} - ")
        run.bold = True
        para.add_run(str(value) if str(value).strip() else "____________________")
    doc.add_paragraph("")
    for line in footer_lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_annex4(v: dict) -> bytes:
    rows = [
        ("Name of Vendor", v["legal_name"]),
        ("Address of Vendor", v["registered_address"]),
        ("Vendor's PAN Number", clean_upper(v["pan"])),
        ("Nature of Service to be provided (for TDS deduction)", v["nature_of_service"]),
        ("Vendor's Contact Person Name", v["contact_person"]),
        ("Vendor's Contact Number", v["contact_mobile"]),
        ("Email Address", v["contact_email"]),
        ("GST Registration Number", clean_upper(v["gstin"])),
        ("Nature of Vendor Firm", v["firm_type"]),
        ("Firm Registration No in ROC", v["firm_reg_no"]),
        ("Linked to Government", v["linked_to_government"]),
        ("MSME Registration", v["msme_registered"]
         + (f" ({v['msme_number']})" if v["msme_registered"] == "Yes" and v["msme_number"] else "")),
        ("Country of Origin", v["country_of_origin"]),
        ("Bank Account Number", v["bank_account_number"]),
        ("Name of Vendor (as per Bank Records)", v["bank_account_name"]),
        ("IFSC / RTGS / NEFT Code", clean_upper(v["ifsc"])),
        ("SWIFT Code (for international transactions)", v["swift"]),
        ("Name of Bank", v["bank_name"]),
        ("Bank-Branch Address", v["bank_branch_address"]),
        ("ESIC Number (if registered)", v["esic_number"]),
    ]
    footer = [
        "",
        "Signed & Stamp: ______________________          Date: ______________",
        "",
        "* Please provide copy of Permanent Account Number, GST Certificate & "
        "Cancelled Cheque for our records.",
        f"(Generated by Hexagon Vendor Registration Agent - Ref {v.get('vendor_id', '')} "
        f"on {date.today().isoformat()})",
    ]
    return _docx_form("Vendor Registration Form (Annex 4)", rows, footer)


def make_206ab(v: dict) -> bytes:
    rows = [
        ("Vendor Code", v.get("vendor_id", "")),
        ("M/s (Legal Name and complete address)",
         f"{v['legal_name']}, {v['registered_address']}"),
        ("PAN", clean_upper(v["pan"])),
        ("TAN", clean_upper(v["tan"]) or "N/A"),
        ("Income Tax Return filed for FY 2023-24 (u/s 139)?",
         v["itr_fy_2023_24"] + (" (acknowledgement to be attached)"
                                if v["itr_fy_2023_24"] == "Yes" else "")),
        ("Income Tax Return filed for FY 2024-25 (u/s 139)?",
         v["itr_fy_2024_25"] + (" (acknowledgement to be attached)"
                                if v["itr_fy_2024_25"] == "Yes" else "")),
        ("Will file Income Tax Return for FY 2025-26 (u/s 139)?",
         v["itr_fy_2025_26_will_file"]),
    ]
    footer = [
        "",
        "I/We hereby declare that the information given above is true and correct, "
        "with reference to the provisions of Section 206AB of the Income Tax Act, 1961.",
        "",
        "Authorised Signatory: ______________________",
        "Name & Designation: ______________________",
        "Place: ______________    Date: ______________    Company Seal:",
        "",
        f"(Generated by Hexagon Vendor Registration Agent - Ref {v.get('vendor_id', '')} "
        f"on {date.today().isoformat()})",
    ]
    return _docx_form("Declaration under Section 206AB of the Income Tax Act, 1961",
                      rows, footer)


def appendix2_draft(v: dict) -> str:
    return (
        "To: financemasterdata.geo@hexagon.com\n"
        "Subject: Vendor bank account approval - " + v["legal_name"] + "\n\n"
        "APPENDIX 2 - Approval request regarding bank account information in a "
        "vendor record (SAP Entities)\n\n"
        "General information\n"
        "SAP entity number: 2580\n"
        "Entity name: Hexagon Geosystems India Pvt Ltd\n"
        f"SAP vendor number: <to be assigned>\n"
        f"Vendor name: {v['legal_name']}\n"
        "New vendor?: Yes\n\n"
        "Bank account information\n"
        f"Account holder name: {v['bank_account_name']}\n"
        f"Account number: {v['bank_account_number']}\n"
        f"IFSC: {clean_upper(v['ifsc'])}\n"
        f"SWIFT: {v['swift'] or 'N/A'}\n"
        f"Bank: {v['bank_name']}, {v['bank_branch_address']}\n\n"
        "Questions\n"
        "How were you informed about this bank account?: Via VRF (Vendor "
        f"Registration Form, ref {v.get('vendor_id', '')})\n"
        "How did you reconfirm the bank account?: Cancelled cheque received"
        + (" and AI-verified against the form" if v.get("ai_verification") else "") + "\n"
    )


# ---------------------------------------------------------------------------
# AI agent (Claude + tools) - conversational registration over the same engine
# ---------------------------------------------------------------------------

_REQUIRED_FOR_PROPOSAL = [
    "legal_name", "registered_address", "firm_type", "contact_person",
    "contact_mobile", "contact_email", "nature_of_service", "service_category",
    "pan", "itr_fy_2023_24", "itr_fy_2024_25", "itr_fy_2025_26_will_file",
    "bank_account_name", "bank_account_number", "ifsc", "bank_name",
]

_PROPOSAL_FIELDS = {
    "legal_name": {"type": "string"},
    "trade_name": {"type": "string"},
    "registered_address": {"type": "string"},
    "firm_type": {"type": "string", "enum": FIRM_TYPES},
    "firm_reg_no": {"type": "string"},
    "linked_to_government": {"type": "string", "enum": ["Yes", "No"]},
    "country_of_origin": {"type": "string"},
    "contact_person": {"type": "string"},
    "contact_mobile": {"type": "string", "description": "10-digit Indian mobile"},
    "contact_email": {"type": "string"},
    "nature_of_service": {"type": "string",
                          "description": "What they will supply/do, for TDS classification"},
    "service_category": {"type": "string", "enum": SERVICE_CATEGORIES},
    "pan": {"type": "string"},
    "tan": {"type": "string"},
    "gstin": {"type": "string"},
    "msme_registered": {"type": "string", "enum": ["Yes", "No"]},
    "msme_number": {"type": "string"},
    "esic_number": {"type": "string"},
    "itr_fy_2023_24": {"type": "string", "enum": ["Yes", "No"]},
    "itr_fy_2024_25": {"type": "string", "enum": ["Yes", "No"]},
    "itr_fy_2025_26_will_file": {"type": "string", "enum": ["Yes", "No"]},
    "bank_account_name": {"type": "string"},
    "bank_account_number": {"type": "string"},
    "ifsc": {"type": "string"},
    "swift": {"type": "string"},
    "bank_name": {"type": "string"},
    "bank_branch_address": {"type": "string"},
}

AGENT_TOOLS = [
    {
        "name": "validate_identifiers",
        "description": ("Validate any Indian statutory identifiers the vendor has "
                        "provided so far: PAN, GSTIN, IFSC, TAN, Udyam number, "
                        "mobile, email. Also cross-checks that the GSTIN embeds "
                        "the PAN. Call this as soon as the user gives any of "
                        "these values - do not wait until the end."),
        "input_schema": {
            "type": "object",
            "properties": {
                "pan": {"type": "string"}, "gstin": {"type": "string"},
                "ifsc": {"type": "string"}, "tan": {"type": "string"},
                "udyam": {"type": "string"}, "mobile": {"type": "string"},
                "email": {"type": "string"},
            },
        },
    },
    {
        "name": "get_registration_status",
        "description": "Look up existing registrations by vendor reference (VND-...) or registered email.",
        "input_schema": {
            "type": "object",
            "properties": {"query": {"type": "string"}},
            "required": ["query"],
        },
    },
    {
        "name": "propose_registration",
        "description": ("Propose the completed vendor registration. Shows the user "
                        "a Confirm button - the system submits only after they "
                        "press it. Call only when every required field is "
                        "collected and identifiers validated clean."),
        "input_schema": {
            "type": "object",
            "properties": _PROPOSAL_FIELDS,
            "required": _REQUIRED_FOR_PROPOSAL,
        },
    },
]

AGENT_SYSTEM = """You are the Hexagon Vendor Registration Agent. You help vendor
representatives register as suppliers to Hexagon Geosystems India, operating
the registration system through your tools.

TODAY is {today}.

WHAT REGISTRATION NEEDS (required): legal name (as per PAN), registered
address, nature of firm ({firm_types}), contact person + 10-digit mobile +
email, nature of service (for TDS classification), service category
({categories}), PAN, the three Section 206AB answers (ITR filed FY 2023-24?
FY 2024-25? will file FY 2025-26?), and bank details (account holder name as
per bank records, account number, IFSC, bank name).
Optional: trade name, ROC/CIN number, GSTIN, TAN, MSME/Udyam number, ESIC,
SWIFT (international only), branch address, country of origin (default INDIA),
linked to government (default No).

RULES:
- Collect details in small batches - at most three questions per message.
- Call validate_identifiers the moment the user provides a PAN, GSTIN, IFSC,
  TAN, Udyam, mobile or email. If something fails, tell them exactly what is
  wrong and ask for a correction before moving on.
- NEVER invent, autocomplete or guess any identifier or bank detail. What the
  user types is what gets validated.
- propose_registration only SHOWS a Confirm button. Never say the registration
  is submitted - say it is ready and ask them to press Confirm. A [SYSTEM
  NOTE] in a user message reports what actually happened - trust it.
- After successful submission the system provides ready-to-sign Annex 4 and
  206AB documents as downloads, and certificates (PAN card, GST certificate,
  cancelled cheque) can be AI-verified in the "New registration" tab.
- The Anti-Corruption Questionnaire (Exhibit 7) must be completed and signed
  by the vendor's authorised representative directly - you never fill it and
  should say so if asked.
- Warm, brief, professional Indian business register. No emojis.
- Only help with vendor registration matters."""


def _agent_system() -> str:
    return AGENT_SYSTEM.format(
        today=date.today().isoformat(),
        firm_types=", ".join(FIRM_TYPES),
        categories=", ".join(SERVICE_CATEGORIES),
    )


def execute_tool(name: str, args: dict) -> dict:
    """Reads execute immediately; the only write path is a pending proposal
    that the user must confirm in the UI."""
    df = load_vendors()

    if name == "validate_identifiers":
        out: dict = {}
        pan = clean_upper(args.get("pan", ""))
        if args.get("pan"):
            out["pan"] = "valid" if PAN_RE.match(pan) else \
                "invalid - expected AAAAA9999A (5 letters, 4 digits, 1 letter)"
        if args.get("gstin"):
            g = clean_upper(args["gstin"])
            if not GSTIN_RE.match(g):
                out["gstin"] = "invalid - expected 15 characters like 06AAAAA9999A1Z5"
            elif pan and PAN_RE.match(pan) and g[2:12] != pan:
                out["gstin"] = (f"format valid BUT it embeds PAN '{g[2:12]}' which does "
                                f"not match the provided PAN '{pan}' - one of them is wrong")
            else:
                out["gstin"] = "valid" + ("" if not pan else " and matches the PAN")
        if args.get("ifsc"):
            out["ifsc"] = "valid" if IFSC_RE.match(clean_upper(args["ifsc"])) else \
                "invalid - expected AAAA0999999 (5th character is zero)"
        if args.get("tan"):
            out["tan"] = "valid" if TAN_RE.match(clean_upper(args["tan"])) else \
                "invalid - expected AAAA99999A"
        if args.get("udyam"):
            out["udyam"] = "valid" if UDYAM_RE.match(str(args["udyam"]).strip().upper()) else \
                "invalid - expected UDYAM-XX-00-0000000"
        if args.get("mobile"):
            out["mobile"] = "valid" if valid_mobile(str(args["mobile"])) else \
                "invalid - must be 10 digits starting 6-9"
        if args.get("email"):
            out["email"] = "valid" if EMAIL_RE.match(str(args["email"]).strip()) else \
                "invalid format"
        return out or {"note": "no identifiers provided to validate"}

    if name == "get_registration_status":
        q = str(args.get("query", "")).strip().upper()
        mine = df[(df["vendor_id"].str.upper() == q)
                  | (df["contact_email"].str.upper() == q)]
        return {"registrations": [
            {"vendor_id": r.vendor_id, "legal_name": r.legal_name,
             "status": r.status, "submitted": r.created_at,
             "note_from_hexagon": r.admin_notes}
            for r in mine.head(5).itertuples()
        ]}

    if name == "propose_registration":
        data = {k: str(args.get(k, "")).strip() for k in _PROPOSAL_FIELDS}
        data["country_of_origin"] = data["country_of_origin"] or "INDIA"
        data["linked_to_government"] = data["linked_to_government"] or "No"
        data["msme_registered"] = data["msme_registered"] or "No"
        problems = validate_vendor(data)
        if problems:
            return {"accepted": False, "problems": problems}
        st.session_state.agent_pending = {"kind": "register", "data": data}
        return {"accepted": True,
                "status": "Proposal shown to the user with a Confirm button. "
                          "Ask them to review it and press Confirm."}

    return {"error": f"Unknown tool: {name}"}


def run_agent(client, api_history: list, display_log: list, max_iters: int = 8):
    retried_empty = False
    for _ in range(max_iters):
        resp = client.messages.create(
            model=CLAUDE_MODEL, max_tokens=3000,
            system=_agent_system(), tools=AGENT_TOOLS, messages=api_history,
        )
        blocks = []
        for b in resp.content:
            if b.type == "text" and b.text.strip():
                blocks.append({"type": "text", "text": b.text})
                display_log.append({"role": "assistant", "kind": "text", "text": b.text})
            elif b.type == "tool_use":
                blocks.append({"type": "tool_use", "id": b.id, "name": b.name,
                               "input": b.input})
                display_log.append({"role": "assistant", "kind": "tool",
                                    "text": f"{b.name} → "
                                            f"{json.dumps(b.input, default=str)[:140]}"})
        if not blocks:  # adaptive-thinking starvation guard
            if retried_empty:
                display_log.append({"role": "assistant", "kind": "text",
                                    "text": "Sorry, I had trouble responding - please try "
                                            "again, or use the New registration tab."})
                return
            retried_empty = True
            time.sleep(1)
            continue
        api_history.append({"role": "assistant", "content": blocks})
        tool_calls = [b for b in blocks if b["type"] == "tool_use"]
        if resp.stop_reason == "tool_use" and tool_calls:
            results = []
            for tc in tool_calls:
                out = execute_tool(tc["name"], tc["input"])
                results.append({"type": "tool_result", "tool_use_id": tc["id"],
                                "content": json.dumps(out, default=str)})
            api_history.append({"role": "user", "content": results})
            continue
        return
    display_log.append({"role": "assistant", "kind": "text",
                        "text": "I've reached my step limit for this request - "
                                "please rephrase or use the form tabs."})


# ---------------------------------------------------------------------------
# Header + admin gate
# ---------------------------------------------------------------------------

def _logo_html() -> str:
    """Embed logo.png from the repo root; fall back to the ⬡ mark."""
    p = Path(__file__).parent / "logo.png"
    if p.exists():
        b64 = base64.b64encode(p.read_bytes()).decode()
        return (f'<img src="data:image/png;base64,{b64}" '
                'style="height:52px;width:auto;" alt="Hexagon"/>')
    return '<span class="hex-title">⬡</span>'

@st.dialog("Admin access")
def admin_login_dialog():
    pw = st.text_input("Admin password", type="password")
    if st.button("Unlock", use_container_width=True):
        if pw and pw == st.secrets.get("ADMIN_PASSWORD", ""):
            st.session_state.admin_ok = True
            st.rerun()
        else:
            st.error("Incorrect password.")


head_l, head_r = st.columns([6, 1])
with head_l:
    st.markdown(
        f'<div style="display:flex;align-items:center;gap:14px;">{_logo_html()}'
        '<span class="hex-title">Hexagon Vendor Registration</span></div>',
        unsafe_allow_html=True)
    st.markdown('<p class="hex-sub">Register once - we prepare your documents. '
                'Annex 4 and the 206AB declaration are generated ready to sign.</p>',
                unsafe_allow_html=True)
with head_r:
    st.write("")
    if st.session_state.get("admin_ok"):
        if st.button("Exit admin", use_container_width=True):
            st.session_state.admin_ok = False
            st.rerun()
    elif st.button("Admin", use_container_width=True):
        admin_login_dialog()

try:
    df_all = load_vendors()
except Exception as e:  # noqa: BLE001
    st.error("Could not connect to the vendor store. Check the Google Sheets "
             f"settings in Streamlit Secrets. ({type(e).__name__})")
    st.stop()

# ---------------------------------------------------------------------------
# Admin dashboard
# ---------------------------------------------------------------------------

if st.session_state.get("admin_ok"):
    st.subheader("Admin - vendor registrations")
    if df_all.empty:
        st.info("No registrations yet.")
        st.stop()

    c1, c2 = st.columns([1, 2])
    status_filter = c1.multiselect("Status", STATUSES, default=STATUSES)
    search = c2.text_input("Search (name / PAN / vendor ID)")
    data = df_all[df_all["status"].isin(status_filter)]
    if search.strip():
        q = search.strip().upper()
        data = data[data["legal_name"].str.upper().str.contains(q, na=False)
                    | data["pan"].str.upper().str.contains(q, na=False)
                    | data["vendor_id"].str.upper().str.contains(q, na=False)]

    st.metric("Registrations shown", len(data))
    st.dataframe(data[["vendor_id", "status", "created_at", "legal_name",
                       "firm_type", "pan", "gstin", "contact_person",
                       "contact_mobile", "service_category"]],
                 use_container_width=True, hide_index=True)

    if not data.empty:
        pick = st.selectbox("Open a registration",
                            [f"{r.vendor_id} · {r.legal_name}" for r in data.itertuples()])
        vid = pick.split(" · ")[0]
        v = data[data["vendor_id"] == vid].iloc[0].to_dict()

        with st.expander("Full details", expanded=True):
            det1, det2 = st.columns(2)
            with det1:
                st.markdown(
                    f"**{v['legal_name']}** ({v['firm_type']})  \n"
                    f"{v['registered_address']}  \n"
                    f"PAN `{v['pan']}` · GSTIN `{v['gstin'] or '-'}` · "
                    f"TAN `{v['tan'] or '-'}`  \n"
                    f"MSME: {v['msme_registered']} {v['msme_number']}  \n"
                    f"Contact: {v['contact_person']} · {v['contact_mobile']} · "
                    f"{v['contact_email']}  \n"
                    f"Service: {v['nature_of_service']} ({v['service_category']})")
            with det2:
                st.markdown(
                    f"**Bank:** {v['bank_name']}  \n{v['bank_branch_address']}  \n"
                    f"A/c `{v['bank_account_number']}` · IFSC `{v['ifsc']}`  \n"
                    f"Holder: {v['bank_account_name']}  \n"
                    f"ITRs: 23-24 {v['itr_fy_2023_24']} · 24-25 {v['itr_fy_2024_25']} · "
                    f"25-26 will file {v['itr_fy_2025_26_will_file']}  \n"
                    f"Docs provided: {v['docs_provided'] or '-'}")
            if v["ai_verification"]:
                st.markdown("**AI verification summary:**")
                st.code(v["ai_verification"], language="json")

        a1, a2, a3 = st.columns(3)
        new_status = a1.selectbox("Set status", STATUSES,
                                  index=STATUSES.index(v["status"])
                                  if v["status"] in STATUSES else 0)
        note = a2.text_input("Admin note", value=v.get("admin_notes", ""))
        a3.write("")
        if a3.button("Save", use_container_width=True):
            if update_vendor(vid, {"status": new_status, "admin_notes": note}):
                st.success("Saved.")
                st.rerun()

        d1, d2, d3 = st.columns(3)
        d1.download_button("Annex 4 (filled .docx)", data=make_annex4(v),
                           file_name=f"{vid}_Annex4.docx", use_container_width=True)
        d2.download_button("206AB Declaration (.docx)", data=make_206ab(v),
                           file_name=f"{vid}_206AB.docx", use_container_width=True)
        with d3.popover("Appendix 2 draft (finance)", use_container_width=True):
            st.caption("Copy into an email to the Finance Master Data Desk:")
            st.code(appendix2_draft(v))
    st.stop()

# ---------------------------------------------------------------------------
# Vendor-facing tabs
# ---------------------------------------------------------------------------

tab_ai, tab_reg, tab_track = st.tabs(
    ["AI Agent", "New registration", "Track my registration"])

with tab_ai:
    st.subheader("Hexagon Vendor Registration Agent")
    st.caption("Register conversationally - the agent collects your details, "
               "validates PAN/GSTIN/IFSC as you go, and prepares the submission. "
               "Nothing is submitted until you press Confirm. Certificate uploads "
               "and AI verification are in the New registration tab.")
    client = get_claude()
    if client is None:
        st.info("The AI agent needs ANTHROPIC_API_KEY in Streamlit Secrets. "
                "Registration via the form tab works without it.")
    else:
        st.session_state.setdefault("agent_api", [])
        st.session_state.setdefault("agent_log", [])
        st.session_state.setdefault("agent_pending", None)
        st.session_state.setdefault("agent_note", "")
        st.session_state.setdefault("agent_last_row", None)

        for entry in st.session_state.agent_log:
            with st.chat_message(entry["role"]):
                if entry.get("kind") == "tool":
                    st.caption(f"🔧 {entry['text']}")
                else:
                    st.markdown(entry["text"])

        if st.session_state.agent_last_row:
            lr = st.session_state.agent_last_row
            g1, g2 = st.columns(2)
            g1.download_button("Annex 4 - filled (.docx)", data=make_annex4(lr),
                               file_name=f"{lr['vendor_id']}_Annex4.docx",
                               use_container_width=True, key="agent_dl_a4")
            g2.download_button("206AB Declaration (.docx)", data=make_206ab(lr),
                               file_name=f"{lr['vendor_id']}_206AB.docx",
                               use_container_width=True, key="agent_dl_206")

        pending = st.session_state.agent_pending
        if pending and pending["kind"] == "register":
            d = pending["data"]
            st.info(f"**Ready to submit:** {d['legal_name']} ({d['firm_type']})  \n"
                    f"PAN `{clean_upper(d['pan'])}` · GSTIN `{clean_upper(d['gstin']) or '-'}`  \n"
                    f"{d['contact_person']} · {d['contact_mobile']} · {d['contact_email']}  \n"
                    f"Bank: {d['bank_name']} · A/c {d['bank_account_number']} · "
                    f"IFSC `{clean_upper(d['ifsc'])}`")
            cc1, cc2 = st.columns(2)
            if cc1.button("Confirm submission", use_container_width=True,
                          key="agent_confirm"):
                fresh = load_vendors()
                problems = validate_vendor(pending["data"])
                if problems:
                    note = "Could not submit: " + " ".join(problems)
                else:
                    row = {c: str(pending["data"].get(c, "")).strip()
                           for c in SHEET_COLUMNS}
                    row["vendor_id"] = gen_vendor_id(fresh)
                    row["status"] = "Submitted"
                    row["pan"] = clean_upper(row["pan"])
                    row["gstin"] = clean_upper(row["gstin"])
                    row["tan"] = clean_upper(row["tan"])
                    row["ifsc"] = clean_upper(row["ifsc"])
                    row["contact_mobile"] = re.sub(r"\D", "", row["contact_mobile"])
                    row["created_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                    row["updated_at"] = row["created_at"]
                    append_vendor(row)
                    st.session_state.agent_last_row = row
                    note = (f"Registration submitted - reference **{row['vendor_id']}**. "
                            "Download the ready-to-sign Annex 4 and 206AB documents "
                            "above, and upload certificates for AI verification in "
                            "the New registration tab if you wish.")
                st.session_state.agent_log.append(
                    {"role": "assistant", "kind": "text", "text": note})
                st.session_state.agent_note = re.sub(r"\*", "", note)
                st.session_state.agent_pending = None
                st.rerun()
            if cc2.button("Discard", use_container_width=True, key="agent_discard"):
                st.session_state.agent_note = "The user discarded the proposal without confirming."
                st.session_state.agent_pending = None
                st.rerun()

        col_reset, _ = st.columns([1, 4])
        if st.session_state.agent_log and col_reset.button("Start over"):
            for k in ("agent_api", "agent_log", "agent_pending",
                      "agent_note", "agent_last_row"):
                st.session_state.pop(k, None)
            st.rerun()

        if prompt := st.chat_input(
                "e.g. I want to register my company as a vendor, or "
                "what's the status of VND-1A2B3C?"):
            content = prompt
            if st.session_state.agent_note:
                content = f"[SYSTEM NOTE: {st.session_state.agent_note}]\n\n{prompt}"
                st.session_state.agent_note = ""
            st.session_state.agent_api.append({"role": "user", "content": content})
            st.session_state.agent_log.append(
                {"role": "user", "kind": "text", "text": prompt})
            with st.spinner("Working…"):
                try:
                    run_agent(client, st.session_state.agent_api,
                              st.session_state.agent_log)
                except Exception as e:  # noqa: BLE001
                    st.session_state.agent_log.append(
                        {"role": "assistant", "kind": "text",
                         "text": f"Something went wrong ({type(e).__name__}). "
                                 "Please try again or use the form tabs."})
            st.rerun()

with tab_reg:
    st.caption("Fields marked * are required. Your details feed every "
               "registration document, so you only enter them once.")

    st.subheader("1 · Company details")
    c1, c2 = st.columns(2)
    with c1:
        legal_name = st.text_input("Legal name (as per PAN) *")
        trade_name = st.text_input("Trade name (if different)")
        firm_type = st.selectbox("Nature of vendor firm *", FIRM_TYPES)
        firm_reg_no = st.text_input("Firm registration no. in ROC / CIN (if applicable)")
    with c2:
        registered_address = st.text_area("Registered address *", height=100)
        linked_gov = st.radio("Linked to Government? *", ["No", "Yes"], horizontal=True)
        country = st.text_input("Country of origin *", value="INDIA")

    st.subheader("2 · Contact & service")
    c3, c4 = st.columns(2)
    with c3:
        contact_person = st.text_input("Contact person name *")
        contact_mobile = st.text_input("Contact mobile *", placeholder="10-digit Indian mobile")
        contact_email = st.text_input("Contact email *")
    with c4:
        service_category = st.selectbox("Service category *", SERVICE_CATEGORIES)
        nature_of_service = st.text_area(
            "Nature of service to be provided (for TDS classification) *", height=100)

    st.subheader("3 · Tax & statutory")
    c5, c6, c7 = st.columns(3)
    with c5:
        pan = st.text_input("PAN *", placeholder="AAAAA9999A")
        tan = st.text_input("TAN (if available)", placeholder="AAAA99999A")
    with c6:
        gstin = st.text_input("GSTIN", placeholder="06AAAAA9999A1Z5")
        esic = st.text_input("ESIC number (if registered)")
    with c7:
        msme = st.radio("MSME / Udyam registered? *", ["No", "Yes"], horizontal=True)
        msme_number = st.text_input("Udyam number", placeholder="UDYAM-HR-06-1234567",
                                    disabled=(msme == "No"))
    st.markdown("**Section 206AB declaration (Income Tax Act)**")
    c8, c9, c10 = st.columns(3)
    itr1 = c8.radio("ITR filed for FY 2023-24? *", ["Yes", "No"], horizontal=True)
    itr2 = c9.radio("ITR filed for FY 2024-25? *", ["Yes", "No"], horizontal=True)
    itr3 = c10.radio("Will file ITR for FY 2025-26? *", ["Yes", "No"], horizontal=True)

    st.subheader("4 · Bank details")
    c11, c12 = st.columns(2)
    with c11:
        bank_account_name = st.text_input("Account holder name (as per bank records) *")
        bank_account_number = st.text_input("Bank account number *")
        ifsc = st.text_input("IFSC *", placeholder="HDFC0001234")
    with c12:
        bank_name = st.text_input("Bank name *")
        bank_branch = st.text_input("Branch address *")
        swift = st.text_input("SWIFT code (international vendors only)")

    st.subheader("5 · Certificates + AI verification")
    st.caption("Upload images or PDFs. The AI reads each certificate and "
               "cross-checks it against what you typed above - mismatches are "
               "flagged before submission. Files are checked in-session; carry "
               "the originals with your signed forms.")
    uploads = {}
    u1, u2 = st.columns(2)
    for i, (key, label) in enumerate(DOC_TYPES.items()):
        target = u1 if i % 2 == 0 else u2
        uploads[key] = target.file_uploader(label, type=["png", "jpg", "jpeg", "pdf"],
                                            key=f"up_{key}")

    vendor_data = {
        "legal_name": legal_name, "trade_name": trade_name,
        "registered_address": registered_address, "firm_type": firm_type,
        "firm_reg_no": firm_reg_no, "linked_to_government": linked_gov,
        "country_of_origin": country,
        "contact_person": contact_person, "contact_mobile": contact_mobile,
        "contact_email": contact_email,
        "nature_of_service": nature_of_service, "service_category": service_category,
        "pan": pan, "tan": tan, "gstin": gstin,
        "msme_registered": msme, "msme_number": msme_number if msme == "Yes" else "",
        "esic_number": esic,
        "itr_fy_2023_24": itr1, "itr_fy_2024_25": itr2,
        "itr_fy_2025_26_will_file": itr3,
        "bank_account_name": bank_account_name,
        "bank_account_number": bank_account_number,
        "ifsc": ifsc, "swift": swift, "bank_name": bank_name,
        "bank_branch_address": bank_branch,
    }

    provided = [k for k, f in uploads.items() if f is not None]
    if provided and st.button("Run AI verification on uploaded certificates"):
        client = get_claude()
        if client is None:
            st.info("AI verification needs ANTHROPIC_API_KEY in secrets - "
                    "you can still submit; documents will be checked manually.")
        else:
            verdicts = {}
            for key in provided:
                f = uploads[key]
                with st.spinner(f"Reading {DOC_TYPES[key]}…"):
                    verdicts[key] = verify_document(
                        client, key, f.getvalue(),
                        f.type or "image/jpeg", vendor_data)
                render_verdict(DOC_TYPES[key], verdicts[key])
            st.session_state.ai_verdicts = verdicts

    st.divider()
    if st.button("Submit registration", use_container_width=True):
        problems = validate_vendor(vendor_data)
        if problems:
            for p in problems:
                st.error(p)
        else:
            row = {c: str(vendor_data.get(c, "")).strip() for c in SHEET_COLUMNS}
            row["vendor_id"] = gen_vendor_id(df_all)
            row["status"] = "Submitted"
            row["pan"] = clean_upper(pan)
            row["gstin"] = clean_upper(gstin)
            row["tan"] = clean_upper(tan)
            row["ifsc"] = clean_upper(ifsc)
            row["contact_mobile"] = re.sub(r"\D", "", contact_mobile)
            row["created_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
            row["updated_at"] = row["created_at"]
            row["docs_provided"] = ", ".join(DOC_TYPES[k] for k in provided)
            verdicts = st.session_state.get("ai_verdicts", {})
            if verdicts:
                row["ai_verification"] = json.dumps(
                    {DOC_TYPES[k]: v for k, v in verdicts.items()},
                    ensure_ascii=False)[:45000]
            append_vendor(row)
            st.success(f"Registration submitted - your reference is "
                       f"**{row['vendor_id']}**. Save it to track your status.")
            st.markdown("**Download, sign, stamp and return these documents "
                        "along with your certificate copies:**")
            g1, g2 = st.columns(2)
            g1.download_button("Annex 4 - Vendor Registration Form (.docx)",
                               data=make_annex4(row),
                               file_name=f"{row['vendor_id']}_Annex4.docx",
                               use_container_width=True)
            g2.download_button("Section 206AB Declaration (.docx)",
                               data=make_206ab(row),
                               file_name=f"{row['vendor_id']}_206AB.docx",
                               use_container_width=True)
            st.caption("Note: the Anti-Corruption Questionnaire (Exhibit 7) is a "
                       "compliance self-declaration and must be completed and "
                       "signed by your authorised representative directly - "
                       "this system deliberately does not fill it for you.")
            st.session_state.pop("ai_verdicts", None)

with tab_track:
    st.subheader("Track my registration")
    ref = st.text_input("Vendor reference (VND-…) or registered email")
    if ref.strip():
        q = ref.strip().upper()
        mine = df_all[(df_all["vendor_id"].str.upper() == q)
                      | (df_all["contact_email"].str.upper() == q)]
        if mine.empty:
            st.info("No registration found for that reference / email.")
        else:
            for r in mine.itertuples():
                badge = {"Submitted": "🟡", "Under Review": "🔵",
                         "Approved": "🟢", "Rejected": "🔴"}.get(r.status, "⚪")
                st.markdown(f"{badge} **{r.vendor_id}** · {r.legal_name}  \n"
                            f"Status: **{r.status}** · Submitted {r.created_at}"
                            + (f"  \nNote from Hexagon: {r.admin_notes}"
                               if r.admin_notes else ""))
